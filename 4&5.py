import requests
from bs4 import BeautifulSoup
from collections import deque
from docx import Document


# Function to extract links from a webpage
# ASSIGNMENT 4
def extract_links(url, soup):
    links = []
    for link in soup.find_all("a", href=True):
        href = link.get("href")
        if href.startswith(("http://", "https://")):
            links.append(href)
        else:
            links.append(url + href)
    return links


# ASSIGNMENT 5
# Crawling only web pages and not documents
def extract_links_modified(url, soup):
    links = []
    skip_extensions = [".jpg", ".jpeg", ".png", ".gif", ".mp4", ".avi", ".pdf"]

    for link in soup.find_all("a", href=True):
        href = link.get("href")

        # Check if the URL starts with 'http://' or 'https://'
        if href.startswith(("http://", "https://")):
            # Check if the URL has any of the skip extensions
            if not any(ext in href.lower() for ext in skip_extensions):
                links.append(href)
        else:
            # Combine the URL and check if it has any of the skip extensions
            full_url = url + href
            if not any(ext in full_url.lower() for ext in skip_extensions):
                links.append(full_url)

    return links


# Starting with the base URL
base_url = "https://pec.edu.in"
visited = set()
queue = deque([base_url])
pages_crawled = 0

# Create a new Word document to store all crawled links
doc = Document()

while pages_crawled < 4 and queue:
    url = queue.popleft()
    if url not in visited:
        try:
            response = requests.get(url)
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, "html.parser")

                # Extract links from the webpage
                links = extract_links_modified(
                    url, soup
                )  # only web pages (assignment 5)

                # Print links found on the current webpage and add to the Word document
                print(f"Links found on {url}:")
                doc.add_paragraph(f"Links found on {url}:")
                for link in links:
                    print(link)
                    doc.add_paragraph(link)

                # Process the webpage here (e.g., extract data)
                # Add new links to the queue
                for link in links:
                    if link not in visited:
                        queue.append(link)

                visited.add(url)
                pages_crawled += 1
                print(f"Crawled {pages_crawled} pages: {url}")
                doc.add_paragraph(f"Crawled {pages_crawled} pages: {url}")
        except Exception as e:
            print(f"Error while crawling {url}: {str(e)}")

# Save the Word document
doc.save("crawled_links.docx")

print("Crawling completed. Results saved in 'crawled_links.docx'.")
